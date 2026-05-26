"""Read Questions_Robert_Mark.docx and print all content as UTF-8."""
import sys
import docx

sys.stdout.reconfigure(encoding="utf-8")

doc = docx.Document("Dependencies/Questions_Robert_Mark.docx")

print("=== PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip()
    if txt:
        print(f"P{i}: {txt[:150]}")

print("\n=== TABLES ===")
for i, table in enumerate(doc.tables):
    print(f"\n--- TABLE {i} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
    for r, row in enumerate(table.rows):
        cells = [c.text.replace("\n", " | ").strip()[:100] for c in row.cells]
        # Deduplicate merged cells (python-docx repeats merged cell text)
        seen = []
        deduped = []
        for c in cells:
            if c not in seen:
                deduped.append(c)
                seen.append(c)
            else:
                deduped.append("(merged)")
        print(f"  R{r:02d}: {deduped}")
