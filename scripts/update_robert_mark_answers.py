"""
Update Dependencies/Questions_Robert_Mark.docx with Robert's confirmed answers
from Recordings 1 & 2 (2026-05-25).

Column layout in Table 3:
  col 0 = Priority
  col 1 = What We Need From You
  col 2 = Your Answer  <-- we fill this
"""
import sys
import shutil
from pathlib import Path
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding="utf-8")

DOC_PATH   = Path("Dependencies/Questions_Robert_Mark.docx")
BACKUP_PATH = Path("Dependencies/Questions_Robert_Mark_ORIGINAL.docx")

# ── back up original before touching it ──────────────────────────────────────
if not BACKUP_PATH.exists():
    shutil.copy(DOC_PATH, BACKUP_PATH)
    print(f"Backup saved → {BACKUP_PATH}")
else:
    print(f"Backup already exists → {BACKUP_PATH}")

# ── answers keyed by row index in Table 3 ────────────────────────────────────
ANSWERS = {
    2: (
        "ANSWERED — 2026-05-25\n\n"
        "Robert completed Recording 1 (AI Quoting & Review Process Guidelines, ~24 min) "
        "and Recording 2 (Quoting & Ordering Workflow, ~14 min) on 2026-05-25. "
        "No walkthrough needed. Both sessions fully transcribed and extracted into "
        "docs/recording_01_ai_quoting_review_guidelines.md and "
        "docs/recording_02_quoting_ordering_workflow.md."
    ),
    4: (
        "ANSWERED — 2026-05-25 (Recording 1)\n\n"
        "Brand-new from-scratch topographic survey: 'Topographic Boundary Survey'\n"
        "Update / topo-only: 'Topo Survey', 'Topographic Survey', or "
        "'Update/Topographic Survey' (any of these three names is used).\n\n"
        "Note: 'Topography Survey' (in our pricing system) maps to these names — "
        "classifier uses the actual FTF service_type field from the order."
    ),
    5: (
        "ANSWERED — 2026-05-25 (Recording 1)\n\n"
        "Construction survey at the design phase maps to 'Topo Survey' in FTF.\n"
        "The DURING-construction sub-services are separate line items: "
        "Building Stake Out, Form Board Survey, Foundation Tie-In.\n\n"
        "These are not 'Construction Survey' — they are individual services "
        "ordered at different stages of the build."
    ),
    6: (
        "ANSWERED — 2026-05-25 (Recording 1)\n\n"
        "A permitting job is still a 'Boundary Survey' in FTF — the difference is "
        "that it requires a 3rd-party digital signature for county portal uploads. "
        "'Permitting' describes the delivery requirement, not a separate service type.\n\n"
        "Classifier rule: service_type = 'Boundary Survey'; flag if notes contain "
        "'Permitting' or 'county portal' to ensure digital signature step is noted."
    ),
    7: (
        "ANSWERED — 2026-05-25 (Recording 1)\n\n"
        "Both names refer to the same service. Robert uses 'Specific Purpose Survey' "
        "as the canonical name in FTF. 'Special Purpose Survey' is also accepted "
        "when it appears in customer requests. Use 'Specific Purpose Survey' "
        "as the standard name in all AI output."
    ),
    9: (
        "PARTIALLY ANSWERED — 2026-05-25 (Recording 2)\n\n"
        "Robert confirmed the list EXISTS and will write it up (referenced as "
        "'Recording 6' in this Q&A doc). A bootstrapped list of 25 competitor names "
        "and 16 domains has been added to config/flag_triggers.py from web research "
        "pending Robert's formal list.\n\n"
        "PENDING: Robert to provide complete list (tracked as I-041). "
        "Current bootstrapped list includes: GT Surveys, Apex Surveying, "
        "Accurate Land Surveyors, Suarez Surveying, Stoner & Associates, "
        "SurvTech Solutions, and others. Robert must validate before go-live."
    ),
    10: (
        "ANSWERED — 2026-05-25 (Recording 1)\n\n"
        "NGE DOES perform:\n"
        "Boundary Survey, Topographic Survey, Topographic Boundary Survey, "
        "Form Board Survey, Spot Survey, Foundation Tie-In Survey, As-Built Survey, "
        "Specific Purpose Survey / Special Purpose Survey, Elevation Certificate, "
        "Plot Plan, Acreage, Elevation Only, Final Survey, Legal Description, "
        "Sketch and Description, Survey Re-draw, Surveyor's Affidavit, "
        "Tree Location, Update Survey.\n\n"
        "NGE does NOT perform:\n"
        "Engineering services / drainage design (NGE is surveying only — no engineers).\n"
        "Site Plans (should be architects or engineers).\n"
        "Wetland Delineation (needs specialist engineer; too complex).\n"
        "Building Stakeout — NGE is 'dabbling' in it again; status ambiguous — "
        "flag all Building Stakeout orders for human review until confirmed back in service (I-042)."
    ),
    11: (
        "ANSWERED — 2026-05-25 (Recording 1)\n\n"
        "Always require human review before sending:\n"
        "• ALTA Table A Survey — management-level review required\n"
        "• B-II Title Review — human review required\n"
        "• Wetland Delineation — never auto-quote; flag immediately\n"
        "• Lot Split — management-level review required\n"
        "• Building Stakeout — flag for human review until confirmed back in service (I-042)\n"
        "• Monroe County (Florida Keys) orders — non-standard pricing, limited crew; "
        "flag for extra review and charge more\n"
        "• Out-of-state properties — flag immediately; NGE is FL-only\n"
        "• Engineering / drainage design requests — reject; NGE does not do engineering\n"
        "• Competitor company name or email domain in order — flag for review\n"
        "• service_type = 'Quote' (unclassified by FTF staff) — hold for manual review"
    ),
    13: (
        "ANSWERED — 2026-05-25 (Recording 1)\n\n"
        "All four listed services require management-level human review before "
        "the AI sends a quote — no exceptions:\n"
        "• ALTA Table A Survey ($1,500) — always human review\n"
        "• B-II Title Review ($450) — always human review\n"
        "• Wetland Delineation ($300) — never auto-quote\n"
        "• Lot Split ($450) — always human review"
    ),
    14: (
        "ANSWERED — 2026-05-25 (Recording 1)\n\n"
        "NGE does NOT perform Wetland Delineation. It requires a specialist engineer "
        "and is too complex for NGE's current capabilities. Added to the "
        "never-auto-quote list in config/flag_triggers.py. Any Wetland Delineation "
        "order must be flagged for human review immediately — do not price or send."
    ),
    15: (
        "ANSWERED — 2026-05-25 (Recording 1)\n\n"
        "NGE operates in Florida only (Phase 1). All 67 FL counties can be quoted.\n\n"
        "Special handling:\n"
        "• Monroe County (Florida Keys): flag for extra review, charge more, "
        "limited crew availability — non-standard pricing applies.\n"
        "• Panhandle / northwest FL (Apalachicola, Tallahassee area): NGE struggles "
        "with crew coverage but still quotes — do not auto-reject.\n"
        "• Vero Beach: having issues but manageable — do not auto-reject.\n"
        "• Strong coverage: Jacksonville, St. Augustine, Orlando, "
        "South / Southeast / Southwest FL.\n\n"
        "Out-of-state: flag immediately and do not quote — NGE is FL-only for Phase 1."
    ),
    17: (
        "ANSWERED — 2026-05-25 (Recording 1)\n\n"
        "Currently NO change order language exists in any NGE estimates. "
        "This is completely new (BRD Amendment 001).\n\n"
        "Robert's process when scope changes mid-project:\n"
        "1. Call or contact the client first\n"
        "2. Explain the scope change\n"
        "3. Get verbal or email confirmation\n"
        "4. THEN add the change order to the invoice\n\n"
        "Change order additions must NEVER be auto-added without explicit "
        "client confirmation. Ryan is drafting the clause text (I-043). "
        "A placeholder clause is live in config/knowledge_base/change_order_clause.txt — "
        "Ryan must review before production go-live."
    ),
    18: (
        "ANSWERED — 2026-05-25 (Recording 1)\n\n"
        "Customer approval methods, most to least common:\n"
        "1. Client pays the invoice via payment link → order auto-advances to 'pending' "
        "(most common — client paying IS the approval)\n"
        "2. Client emails confirmation: 'please proceed', 'we accept', etc.\n"
        "3. Client accepts via the FTF portal\n"
        "4. Phone call — Robert always asks for an email follow-up to document the approval"
    ),
    19: (
        "PENDING — not yet provided\n\n"
        "Robert did not address past dispute history in either recording. "
        "This is a STANDARD priority item — Ryan needs this context to write "
        "the change order clause that covers the most common friction points. "
        "Please ask Robert to provide 2-3 examples if any disputes have occurred."
    ),
}

GENERAL_NOTES_ANSWER = (
    "Additional context confirmed by Robert (Recording 2, 2026-05-25):\n\n"
    "SUMMIT'S ROLE: Robert is NOT normally the one creating orders in FTF. Summit posts "
    "suggested prices in Microsoft Teams ('Blue Invoicing' and 'Yellow Invoicing' channels). "
    "Robert, Alan, and Mark review Summit's price, check the GIS map, check client history, "
    "then confirm or adjust — THEN the quote is sent. This is a deliberate review step, "
    "not a rubber-stamp.\n\n"
    "PRICING FACTORS (how Robert decides a price):\n"
    "1. Client sales history (what price they've accepted before)\n"
    "2. Property features (pool, seawall, canal, right-of-ways — checked on GIS map)\n"
    "3. Area / county / local market\n"
    "4. Platted vs. unplatted (affects survey complexity)\n"
    "5. Scope of work\n"
    "6. Competitive positioning vs. other FL surveyors for this client\n\n"
    "QUOTE EXPIRY: FTF auto-cancels any quote older than 60 days. AI pipeline must "
    "prioritise sending within this window.\n\n"
    "'SMITH ZONE' FLAG: Track Flow shows a yellow-highlighted territory called the "
    "'Smith Zone'. Details and pricing rules unknown — to be confirmed by Robert (I-048).\n\n"
    "WORKFLOW CHANNELS: Quote requests come via (1) Outlook email, "
    "(2) Microsoft Teams 'Blue Invoicing' channel, (3) nexgensurveying.com web form. "
    "Robert uses county property appraiser GIS (bcpa.net for Broward, pbcgov.com for "
    "Palm Beach) to visually confirm parcel, lot size, and flood zone before pricing."
)


def set_cell_text(cell, text: str) -> None:
    """Clear cell and write new text, preserving basic paragraph structure."""
    # Remove all existing paragraphs except the first (python-docx requires at least one)
    for para in cell.paragraphs[1:]:
        p = para._element
        p.getparent().remove(p)

    first_para = cell.paragraphs[0]
    first_para.clear()

    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            run = first_para.add_run(line)
        else:
            para = cell.add_paragraph()
            run = para.add_run(line)
        # Bold the first line (status / answered line)
        if i == 0 and (line.startswith("ANSWERED") or line.startswith("PARTIALLY") or line.startswith("PENDING")):
            run.bold = True


# ── load and update ───────────────────────────────────────────────────────────
doc = docx.Document(DOC_PATH)
table3 = doc.tables[3]  # The main Q&A table

for row_idx, answer_text in ANSWERS.items():
    cell = table3.rows[row_idx].cells[2]
    set_cell_text(cell, answer_text)
    print(f"  Updated R{row_idx:02d} — {answer_text[:60]}...")

# Update General Notes table (Table 4, row 0, col 1)
table4 = doc.tables[4]
set_cell_text(table4.rows[0].cells[1], GENERAL_NOTES_ANSWER)
print("  Updated General Notes section")

doc.save(DOC_PATH)
print(f"\nSaved → {DOC_PATH}")
print("Done.")
